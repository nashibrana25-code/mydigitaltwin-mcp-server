import re
from pathlib import Path
from docx import Document

MD_FILE = Path('WEEK7_DELIVERABLE_SUBMISSION.md')
OUT_FILE = Path('WEEK7_DELIVERABLE_SUBMISSION.docx')

HEADING_PATTERNS = [
    (re.compile(r'^#{1}\s+(.*)'), 0),
    (re.compile(r'^#{2}\s+(.*)'), 1),
    (re.compile(r'^#{3}\s+(.*)'), 2),
    (re.compile(r'^#{4}\s+(.*)'), 3),
]

BULLET_PATTERN = re.compile(r'^[-*]\s+(.*)')
TABLE_SEPARATOR = re.compile(r'^\|')


def add_table(lines, doc):
    """Robust Markdown table ingestion.
    Accepts a list of raw markdown table lines including header & separator.
    Automatically pads/truncates irregular rows to match header length.
    Falls back to plain paragraphs if header parse fails.
    """
    if len(lines) < 2:
        # Not enough lines to form a table; dump as paragraphs
        for ln in lines:
            doc.add_paragraph(ln)
        return
    header_raw = lines[0].strip('|')
    header = [h.strip() for h in header_raw.split('|') if h.strip()]
    if not header:
        for ln in lines:
            doc.add_paragraph(ln)
        return
    col_count = len(header)
    table = doc.add_table(rows=1, cols=col_count)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
    # Process data rows (skip separator line index 1)
    for ln in lines[2:]:
        if not ln.startswith('|'):
            break
        row_cells_raw = [c.strip() for c in ln.strip('|').split('|')]
        # Normalize row length
        if len(row_cells_raw) < col_count:
            row_cells_raw.extend([''] * (col_count - len(row_cells_raw)))
        elif len(row_cells_raw) > col_count:
            row_cells_raw = row_cells_raw[:col_count]
        row_cells = table.add_row().cells
        for i, c in enumerate(row_cells_raw):
            row_cells[i].text = c


def main():
    if not MD_FILE.exists():
        raise SystemExit(f"Markdown file {MD_FILE} not found")
    text = MD_FILE.read_text(encoding='utf-8').splitlines()
    doc = Document()

    table_buffer = []
    in_table = False

    for line in text:
        if in_table:
            if line.strip() == '' or not line.startswith('|'):
                # finalize table
                add_table(table_buffer, doc)
                table_buffer.clear()
                in_table = False
            else:
                table_buffer.append(line)
                continue
        if line.startswith('|'):
            in_table = True
            table_buffer.append(line)
            continue
        if not line.strip():
            doc.add_paragraph('')
            continue
        # Headings
        matched = False
        for pattern, level in HEADING_PATTERNS:
            m = pattern.match(line)
            if m:
                text_head = m.group(1).strip()
                if level == 0:
                    doc.add_heading(text_head, level=0)
                else:
                    doc.add_heading(text_head, level=level+1)
                matched = True
                break
        if matched:
            continue
        # Bullets
        b = BULLET_PATTERN.match(line)
        if b:
            doc.add_paragraph(b.group(1), style='List Bullet')
            continue
        # Fallback paragraph
        doc.add_paragraph(line)

    if table_buffer:
        add_table(table_buffer, doc)

    doc.save(OUT_FILE)
    print(f"Saved {OUT_FILE}")

if __name__ == '__main__':
    main()